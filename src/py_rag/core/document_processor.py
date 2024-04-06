import asyncio
import hashlib
import logging
from datetime import datetime
from pathlib import Path
from typing import Union, List, Dict, Any

import aiofiles
import chardet
import fitz
import pandas as pd
from docx import Document
from openpyxl import load_workbook

from py_faiss.config import settings

logger = logging.getLogger(__name__)


async def _extract_from_docx(file_path: Path) -> str:
    """从 DOCX 文件提取文本"""
    try:
        # 在线程池中执行 IO 密集型操作
        loop = asyncio.get_event_loop()

        def _read_docx():
            doc = Document(file_path)
            paragraphs = []

            # 提取段落文件
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)

            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(' | '.join(row_text))

            return '\n'.join(paragraphs)

        text = await loop.run_in_executor(None, _read_docx)
        return text

    except Exception as e:
        raise Exception(f"DOCX 处理失败: {e}")


async def _extract_from_pdf(file_path: Path) -> str:
    """从 PDF 文件提取文本"""
    try:
        loop = asyncio.get_event_loop()

        def _read_pdf():
            text_content = []

            doc = fitz.open(file_path)

            for page_num in range(doc.page_count):
                page = doc[page_num]
                page_text = page.get_text("text")
                if page_text.strip():
                    text_content.append(f"[页面 {page_num + 1}]\n{page_text.strip()}")

            return '\n\n'.join(text_content)

        text = await loop.run_in_executor(None, _read_pdf)
        return text

    except Exception as e:
        raise Exception(f"PDF 处理失败: {e}")


async def _extract_from_text(file_path: Path) -> str:
    """从文本文件提取内容"""
    try:
        # 检测文件编码
        async with aiofiles.open(file_path, 'rb') as f:
            raw_data = await f.read()
            encoding_result = chardet.detect(raw_data)
            encoding = encoding_result['encoding'] or 'utf-8'

        # 读取文本
        async with aiofiles.open(file_path, 'r', encoding=encoding) as f:
            text = await f.read()
            return text.strip()

    except Exception as e:
        # 尝试常见编码
        for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
            try:
                async with aiofiles.open(file_path, 'r', encoding=encoding) as f:
                    text = await f.read()
                    return text.strip()
            except Exception as e:
                logger.warning(f"尝试编码 {encoding} 失败: {e}")
                continue

        raise Exception(f"文本文件读取失败: {e}")


async def _extract_from_excel(file_path: Path) -> str:
    """从 Excel 文件提取文本"""
    try:
        loop = asyncio.get_event_loop()

        def _read_excel():
            content = []

            if file_path.suffix.lower() == '.xlsx':
                # 使用 openpyxl 读取 .xlsx
                wb = load_workbook(file_path, read_only=True)
                for sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]
                    content.append(f"[工作表: {sheet_name}]")

                    for row in sheet.iter_rows(values_only=True):
                        row_data = [str(cell) if cell is not None else '' for cell in row]
                        row_text = ' | '.join(filter(None, row_data))
                        if row_text.strip():
                            content.append(row_text)
                    content.append("")
            else:
                # 使用 pandas 读取 .xls
                excel_file = pd.ExcelFile(file_path)
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    content.append(f"[工作表: {sheet_name}]")

                    # 添加列标题
                    headers = ' | '.join(str(col) for col in df.columns)
                    content.append(headers)

                    # 添加数据行
                    for _, row in df.iterrows():
                        row_data = [str(val) if pd.notna(val) else '' for val in row]
                        row_text = ' | '.join(filter(None, row_data))
                        if row_text.strip():
                            content.append(row_text)
                    content.append("")

            return '\n'.join(content)

        text = await loop.run_in_executor(None, _read_excel)
        return text

    except Exception as e:
        raise Exception(f"Excel 处理失败: {e}")


async def _extract_from_csv(file_path: Path) -> str:
    """从 CSV 文件提取文本"""
    try:
        loop = asyncio.get_event_loop()

        def _read_csv():
            # 尝试不同编码
            for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    break
                except:
                    continue
            else:
                raise Exception("无法识别 CSV 文件编码")

            content = []

            # 添加列标题
            headers = ' | '.join(str(col) for col in df.columns)
            content.append(headers)

            # 添加数据行
            for _, row in df.iterrows():
                row_data = [str(val) if pd.notna(val) else '' for val in row]
                row_text = ' | '.join(filter(None, row_data))
                if row_text.strip():
                    content.append(row_text)

            return '\n'.join(content)

        text = await loop.run_in_executor(None, _read_csv)
        return text

    except Exception as e:
        raise Exception(f"CSV 处理失败: {e}")


async def _extract_from_json(file_path: Path) -> str:
    """从 JSON 文件提取文本"""
    try:
        import json

        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            content = await f.read()
            data = json.loads(content)

            # 递归提取 JSON 中的文本内容
            def extract_text_from_json(obj, path=""):
                texts = []

                if isinstance(obj, dict):
                    for key, value in obj.items():
                        current_path = f"{path}.{key}" if path else key
                        texts.extend(extract_text_from_json(value, current_path))
                elif isinstance(obj, list):
                    for i, item in enumerate(obj):
                        current_path = f"{path}[{i}]" if path else f"[{i}]"
                        texts.extend(extract_text_from_json(item, current_path))
                else:
                    # 基本类型，转换为字符串
                    text = str(obj).strip()
                    if text and len(text) > 1:  # 过滤单字符
                        texts.append(f"{path}: {text}")

                return texts

            text_parts = extract_text_from_json(data)
            return '\n'.join(text_parts)

    except Exception as e:
        raise Exception(f"JSON 处理失败: {e}")


async def _extract_from_xml(file_path: Path) -> str:
    """从 XML 文件提取文本"""
    try:
        import xml.etree.ElementTree as ET

        loop = asyncio.get_event_loop()

        def _read_xml():
            tree = ET.parse(file_path)
            root = tree.getroot()

            def extract_xml_text(element, level=0):
                texts = []
                indent = "  " * level

                # 元素标签和属性
                if element.attrib:
                    attrs = ", ".join(f"{k}={v}" for k, v in element.attrib.items())
                    texts.append(f"{indent}<{element.tag} {attrs}>")
                else:
                    texts.append(f"{indent}<{element.tag}>")

                # 元素文本内容
                if element.text and element.text.strip():
                    texts.append(f"{indent}  {element.text.strip()}")

                # 递归处理子元素
                for child in element:
                    texts.extend(extract_xml_text(child, level + 1))

                return texts

            text_parts = extract_xml_text(root)
            return '\n'.join(text_parts)

        text = await loop.run_in_executor(None, _read_xml)
        return text

    except Exception as e:
        raise Exception(f"XML 处理失败: {e}")


class DocumentProcessor:
    """文档处理器 - 支持多种文档格式的文本提取和处理"""

    def __init__(self):
        self.supported_extensions = {
            '.docx', '.doc',  # Word 文档
            '.pdf',  # PDF 文档
            '.txt', '.md',  # 文本文档
            '.xlsx', '.xls',  # Excel 文档
            '.csv',  # CSV 文档
            '.json',  # JSON 文档
            '.xml',  # XML 文档
        }

        # 文本分割配置
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP

        # 临时文件目录
        self.temp_dir = Path(settings.TEMP_PATH)
        self.temp_dir.mkdir(parents=True, exist_ok=True)

    async def extract_text(self, file_path: Union[str, Path]):
        """
        从文档中提取文本

        Args:
            file_path: 文档文件路径

        Returns:
            提取的文本内容
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        # 检查文件大小
        file_size = file_path.stat().st_size
        if file_size > settings.MAX_FILE_SIZE:
            raise ValueError(f"文件过大: {file_size / 1024 / 1024:.1f}MB > {settings.MAX_FILE_SIZE / 1024 / 1024}MB")

        # 获取文件扩展名
        extension = file_path.suffix.lower()

        if extension not in self.supported_extensions:
            raise ValueError(f"不支持的文件格式: {extension}")

        try:
            # 根据文件类型选择提取方法
            if extension in ['.docx', '.doc']:
                text = await _extract_from_docx(file_path)
            elif extension == '.pdf':
                text = await _extract_from_pdf(file_path)
            elif extension in ['.txt', '.md']:
                text = await _extract_from_text(file_path)
            elif extension in ['.xlsx', '.xls']:
                text = await _extract_from_excel(file_path)
            elif extension == '.csv':
                text = await _extract_from_csv(file_path)
            elif extension == '.json':
                text = await _extract_from_json(file_path)
            elif extension == '.xml':
                text = await _extract_from_xml(file_path)
            else:
                raise ValueError(f"暂不支持的文件格式: {extension}")

            logger.info(f"✅ 成功提取文本: {file_path.name}, 长度: {len(text)} 字符")
            return text

        except Exception as e:
            logger.error(f"❌ 文本提取失败 {file_path.name}: {e}")
            raise

    def split_text(self, text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
        """
        将文本分割成块

        Args:
            text: 输入文本
            chunk_size: 块大小，默认使用配置值
            overlap: 重叠大小，默认使用配置值

        Returns:
            文本块列表
        """
        if not text or not text.strip():
            return []

        chunk_size = chunk_size or self.chunk_size
        overlap = overlap or self.chunk_overlap

        # 确保 overlap 不大于 chunk_size
        overlap = min(overlap, chunk_size // 2)

        chunks = []
        start = 0
        text_length = len(text)

        while start < text_length:
            end = min(start + chunk_size, text_length)
            chunk = text[start:end].strip()

            if chunk:
                chunks.append(chunk)

            # 如果已经到达文本末尾，退出循环
            if end >= text_length:
                break

            # 计算下一个块的起始位置
            start = end - overlap

            # 避免无限循环
            if start <= chunks.__len__() * chunk_size - overlap * chunks.__len__():
                start = end

        return chunks

    def smart_split_text(self, text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
        """
        智能文本分割 - 尽量在句子边界分割

        Args:
            text: 输入文本
            chunk_size: 块大小
            overlap: 重叠大小

        Returns:
            文本块列表
        """
        if not text or not text.strip():
            return []

        chunk_size = chunk_size or self.chunk_size
        overlap = overlap or self.chunk_overlap

        # 句子分隔符
        sentence_separators = ['.', '!', '?', '。', '！', '？', '\n\n']

        chunks = []
        current_chunk = ""
        sentences = []

        # 简单的句子分割
        temp_text = text
        for sep in sentence_separators:
            temp_text = temp_text.replace(sep, sep + '<SPLIT>')

        parts = temp_text.split('<SPLIT>')

        for part in parts:
            part = part.strip()
            if not part:
                continue

            # 如果当前块加上新句子不超过限制，则添加
            if len(current_chunk) + len(part) <= chunk_size:
                current_chunk += part
            else:
                # 保存当前块
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())

                # 如果单个句子太长，需要强制分割
                if len(part) > chunk_size:
                    # 使用基本分割方法
                    sub_chunks = self.split_text(part, chunk_size, overlap)
                    chunks.extend(sub_chunks[:-1])  # 除了最后一个
                    current_chunk = sub_chunks[-1] if sub_chunks else ""
                else:
                    current_chunk = part

        # 添加最后一个块
        if current_chunk.strip():
            chunks.append(current_chunk.strip())

        # 应用重叠逻辑
        if overlap > 0 and len(chunks) > 1:
            overlapped_chunks = []
            for i, chunk in enumerate(chunks):
                if i == 0:
                    overlapped_chunks.append(chunk)
                else:
                    # 从前一个块取 overlap 长度的文本
                    prev_chunk = chunks[i - 1]
                    overlap_text = prev_chunk[-overlap:] if len(prev_chunk) > overlap else prev_chunk
                    overlapped_chunk = overlap_text + " " + chunk
                    overlapped_chunks.append(overlapped_chunk)

            return overlapped_chunks

        return chunks

    async def process_document(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        完整处理文档：提取文本 + 分割

        Args:
            file_path: 文档路径

        Returns:
            处理结果字典
        """
        file_path = Path(file_path)
        start_time = datetime.now()

        try:
            # 提取文本
            text = await self.extract_text(file_path)

            # 分割文本
            chunks = self.smart_split_text(text)

            # 计算文档指纹
            document_hash = hashlib.md5(text.encode()).hexdigest()

            # 处理时间
            processing_time = (datetime.now() - start_time).total_seconds()

            result = {
                'file_path': str(file_path),
                'file_name': file_path.name,
                'file_size': file_path.stat().st_size,
                'document_hash': document_hash,
                'text_length': len(text),
                'chunks_count': len(chunks),
                'chunks': chunks,
                'processing_time': processing_time,
                'processed_at': datetime.now().isoformat(),
                'status': 'success'
            }

            logger.info(f"文档处理完成: {file_path.name}, "
                        f"文本长度: {len(text)}, 块数: {len(chunks)}, "
                        f"耗时: {processing_time:.2f}s")

            return result

        except Exception as e:
            error_result = {
                'file_path': str(file_path),
                'file_name': file_path.name,
                'status': 'error',
                'error': str(e),
                'processed_at': datetime.now().isoformat(),
                'processing_time': (datetime.now() - start_time).total_seconds()
            }

            logger.error(f"❌ 文档处理失败: {file_path.name}, 错误: {e}")
            return error_result

    async def save_temp_file(self, content: bytes, filename: str) -> Path:
        """
        保存临时文件

        Args:
            content: 文件内容
            filename: 文件名

        Returns:
            临时文件路径
        """
        # 生成唯一文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        temp_filename = f"{timestamp}_{filename}"
        temp_path = self.temp_dir / temp_filename

        async with aiofiles.open(temp_path, 'wb') as f:
            await f.write(content)

        logger.info(f"✅ 临时文件已保存: {temp_path}")
        return temp_path

    async def cleanup_temp_files(self, max_age_hours: int = 24):
        """
        清理临时文件

        Args:
            max_age_hours: 文件最大保留时间（小时）
        """
        current_time = datetime.now()
        cleaned_count = 0

        for file_path in self.temp_dir.iterdir():
            if file_path.is_file():
                file_age = current_time - datetime.fromtimestamp(file_path.stat().st_mtime)

                if file_age.total_seconds() > max_age_hours * 3600:
                    try:
                        file_path.unlink()
                        cleaned_count += 1
                        logger.debug(f"已删除临时文件: {file_path}")
                    except Exception as e:
                        logger.warning(f"删除临时文件失败 {file_path}: {e}")

        if cleaned_count > 0:
            logger.info(f"清理了 {cleaned_count} 个临时文件")

    def get_supported_types(self) -> List[str]:
        """获取支持的文件类型列表"""
        return list(self.supported_extensions)

    def is_supported_file(self, file_path: Union[str, Path]) -> bool:
        """检查文件是否被支持"""
        file_path = Path(file_path)
        return file_path.suffix.lower() in self.supported_extensions


# 全局实例
document_processor = DocumentProcessor()
